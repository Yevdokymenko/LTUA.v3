import streamlit as st
import os
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from docx import Document
import re

from translate_script import (
    extract_text,
    translate_text_google,
    chunk_paragraphs,
    translate_chunk_openai,
    create_translation_table,
    setup_document_orientation,
    add_title,
    sanitize_filename # Переконайтесь, що ця функція є у translate_script.py або визначена тут
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")

st.set_page_config(page_title="LegalTransUA", layout="wide")

TEMP_DIR = "temp"
if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

def save_uploaded_file(uploaded_file):
    file_path = os.path.join(TEMP_DIR, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path

st.image("https://i.imgur.com/JmLIg6y.jpeg", width='stretch')

# --- Функція для перевірки пароля ---
def check_password():
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    if not st.session_state["password_correct"]:
        st.title("Вхід до системи")
        password = st.text_input("Введіть пароль для доступу:", type="password")
        if st.button("Увійти"):
            if password == "170287":
                st.session_state["password_correct"] = True
                st.rerun()
            else:
                st.error("Пароль невірний.")
        return False
    else:
        return True

# --- Основний блок програми ---
if check_password():
    st.sidebar.title("Навігація")
    section = st.sidebar.radio(
        "Оберіть розділ:",
        ["Переклад документів", "Про додаток", "Корисні посилання", "Допомога Україні", "Контакти"]
    )

    if section == "Переклад документів":
        st.title("Система порівняльного перекладу документів")
        st.header("Переклад за допомогою Google Translate та OpenAI")
        st.write("Завантажте документ у форматі .docx або .pdf, або надайте посилання на веб-сторінку для одночасного перекладу двома системами.")

        type_of_source = st.radio("Оберіть тип джерела:", ["Завантажити файл", "Вказати URL"])
        chunk_size = 5

        if type_of_source == "Завантажити файл":
            uploaded_file = st.file_uploader("Виберіть файл (DOCX або PDF):", type=["docx", "pdf"])
            if uploaded_file:
                file_path = save_uploaded_file(uploaded_file)
                st.success(f"Файл '{uploaded_file.name}' успішно завантажено.")

                if st.button("Розпочати переклад"):
                    # ... (вся логіка перекладу залишається незмінною)
                    paragraphs = extract_text(file_path)
                    if not paragraphs:
                        st.warning("Не вдалося знайти текст у документі.")
                    else:
                        st.info(f"Знайдено {len(paragraphs)} абзаців для перекладу.")

                        google_progress = st.progress(0, text="Переклад Google Translate...")
                        google_trans = ["" for _ in paragraphs]
                        with ThreadPoolExecutor(max_workers=5) as executor:
                            futures_g = {executor.submit(translate_text_google, p): i for i, p in enumerate(paragraphs)}
                            done_g = 0
                            for future in as_completed(futures_g):
                                idx = futures_g[future]
                                google_trans[idx] = future.result()
                                done_g += 1
                                frac_g = done_g / len(paragraphs)
                                google_progress.progress(frac_g, text=f"Google Translate: {int(frac_g*100)}%")

                        openai_progress = st.progress(0, text="Переклад OpenAI...")
                        openai_trans = [None] * len(paragraphs)
                        all_chunks = list(chunk_paragraphs(paragraphs, chunk_size=chunk_size))
                        total_chunks = len(all_chunks)
                        for chunk_i, chunk in enumerate(all_chunks):
                            chunk_result = translate_chunk_openai(chunk)
                            chunk_start = chunk_i * chunk_size
                            for j, translation in enumerate(chunk_result):
                                openai_trans[chunk_start + j] = translation
                            frac_o = (chunk_i + 1) / total_chunks
                            openai_progress.progress(frac_o, text=f"OpenAI: {int(frac_o*100)}%")

                        doc = Document()
                        setup_document_orientation(doc)
                        add_title(doc)
                        doc.add_paragraph(f"Дата та час перекладу: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        create_translation_table(doc, paragraphs, google_trans, openai_trans)

                        timestamp_str = datetime.now().strftime("%Y-%m-%d")
                        base_name = os.path.splitext(uploaded_file.name)[0]
                        base_name = sanitize_filename(base_name)
                        output_filename = f"{base_name} (Translated by LTUA {timestamp_str}).docx"
                        output_path = os.path.join(TEMP_DIR, output_filename)
                        doc.save(output_path)

                        st.success("Переклад успішно завершено!")
                        with open(output_path, "rb") as file:
                            st.download_button(
                                label="Завантажити результат (.docx)",
                                data=file,
                                file_name=output_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

        elif type_of_source == "Вказати URL":
            url = st.text_input("Введіть посилання на веб-сторінку:")
            if url and st.button("Розпочати переклад"):
                st.info(f"Аналізуємо сторінку за посиланням...")
                try:
                    paragraphs = extract_text(url)
                except Exception as e:
                    st.error(f"Не вдалося обробити посилання: {e}")
                    paragraphs = []

                if not paragraphs:
                    st.warning("На сторінці не знайдено абзаців тексту для перекладу.")
                else:
                    st.success(f"Знайдено {len(paragraphs)} абзаців для перекладу.")
                    # ... (тут та сама логіка перекладу, що і для файлу)
                    google_progress = st.progress(0, text="Переклад Google Translate...")
                    google_trans = ["" for _ in paragraphs]
                    with ThreadPoolExecutor(max_workers=5) as executor:
                        futures_g = {executor.submit(translate_text_google, p): i for i, p in enumerate(paragraphs)}
                        done_g = 0
                        for future in as_completed(futures_g):
                            idx = futures_g[future]
                            google_trans[idx] = future.result()
                            done_g += 1
                            frac_g = done_g / len(paragraphs)
                            google_progress.progress(frac_g, text=f"Google Translate: {int(frac_g*100)}%")

                    openai_progress = st.progress(0, text="Переклад OpenAI...")
                    openai_trans = [None] * len(paragraphs)
                    all_chunks = list(chunk_paragraphs(paragraphs, chunk_size=chunk_size))
                    total_chunks = len(all_chunks)
                    for chunk_i, chunk in enumerate(all_chunks):
                        chunk_result = translate_chunk_openai(chunk)
                        chunk_start = chunk_i * chunk_size
                        for j, translation in enumerate(chunk_result):
                            openai_trans[chunk_start + j] = translation
                        frac_o = (chunk_i + 1) / total_chunks
                        openai_progress.progress(frac_o, text=f"OpenAI: {int(frac_o*100)}%")

                    doc = Document()
                    setup_document_orientation(doc)
                    add_title(doc)
                    doc.add_paragraph(f"Дата та час перекладу: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    create_translation_table(doc, paragraphs, google_trans, openai_trans)

                    timestamp_str = datetime.now().strftime("%Y-%m-%d")
                    output_filename = f"Web-document (Translated by LTUA {timestamp_str}).docx"
                    output_path = os.path.join(TEMP_DIR, output_filename)
                    doc.save(output_path)

                    st.success("Переклад успішно завершено!")
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label="Завантажити результат (.docx)",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

    elif section == "Про додаток":
        st.title("Про LegalTransUA")
        st.write("""
        **LegalTransUA** — це інструмент для порівняльного перекладу юридичних та інших текстів.
        Він дозволяє одночасно отримати результати від двох провідних систем машинного перекладу:
        **Google Translate** та **OpenAI GPT**. Це дає змогу юристам, перекладачам та іншим фахівцям
        швидко порівнювати варіанти перекладу та обирати найкращий.

        Додаток підтримує роботу з файлами `.docx`, `.pdf`, а також аналіз тексту з веб-сторінок.
        """)

    elif section == "Корисні посилання":
        st.title("Корисні посилання")
        st.markdown("""
        - [EUR-Lex: Доступ до законодавства Європейського Союзу](https://eur-lex.europa.eu/)
        - [Законодавство України: Портал Верховної Ради](https://zakon.rada.gov.ua/)
        - [Linguee: Контекстний словник для перекладачів](https://www.linguee.com/)
        - [ProZ.com: Спільнота професійних перекладачів](https://www.proz.com/)
        """)

    elif section == "Допомога Україні":
        st.title("Підтримайте Україну")
        st.write("""
        Ваша підтримка є надзвичайно важливою. Ви можете зробити внесок у перевірені фонди,
        які допомагають Збройним Силам України та гуманітарним ініціативам.
        """)
        st.markdown("""
        - [Фонд "Повернись живим"](https://savelife.in.ua/)
        - [UNITED24: Офіційна фандрейзингова платформа України](https://u24.gov.ua/)
        """)

    elif section == "Контакти":
        st.title("Контакти")
        st.write("З питань співпраці, пропозицій або технічних проблем:")
        st.markdown("""
        - **Email:** yevdokymenkodn@gmail.com
        - **LinkedIn:** [Dmytro Yevdokymenko](https://www.linkedin.com/in/yevdokymenko/)
        """)