import logging
import os
import re
import time
import requests
import shutil
import fitz  # PyMuPDF
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator

import openai
import streamlit as st

# Налаштування логування
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# --- ІНІЦІАЛІЗАЦІЯ КЛІЄНТА OPENROUTER ---
# Використовуємо ключ із Secrets під назвою OPENAI_API_KEY, але стукаємо на адресу OpenRouter
try:
    api_key = st.secrets["OPENAI_API_KEY"]
except Exception:
    # Fallback для локального тестування, якщо st.secrets недоступні
    api_key = os.getenv("OPENAI_API_KEY", "YOUR_KEY_HERE")

client = openai.OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=api_key
)

# --- ДОПОМІЖНІ ФУНКЦІЇ ДЛЯ ТЕКСТУ ---

def extract_text_from_docx(file_path: str):
    """Витягує текст із DOCX-файлу, повертає список абзаців."""
    doc = docx.Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]


def extract_text_from_pdf(file_path: str):
    """Витягує текст із PDF-файлу, повертає список абзаців."""
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text("text") + "\n"
    return [line.strip() for line in full_text.splitlines() if line.strip()]


def extract_text_from_html(url: str):
    """Екстрагує текст із веб-сторінки, повертає список абзаців."""
    response = requests.get(url)
    if response.status_code != 200:
        raise Exception(f"Не вдалося завантажити сторінку: {url}")

    soup = BeautifulSoup(response.content, "html.parser")
    paragraphs = soup.find_all("p")
    return [para.get_text().strip() for para in paragraphs if para.get_text().strip()]


def extract_text(source: str):
    """Визначає тип джерела (PDF, DOCX, URL) і повертає список абзаців."""
    if source.startswith("http"):
        logging.info("Джерело визначено як веб-сторінка.")
        return extract_text_from_html(source)
    elif source.endswith(".pdf"):
        logging.info("Джерело визначено як PDF-файл.")
        return extract_text_from_pdf(source)
    elif source.endswith(".docx"):
        logging.info("Джерело визначено як DOCX-файл.")
        return extract_text_from_docx(source)
    else:
        raise ValueError("Формат файлу не підтримується. Підтримуються DOCX, PDF або URL.")


def sanitize_filename(filename: str) -> str:
    """Очищає ім'я файлу від недопустимих символів."""
    name, ext = os.path.splitext(filename)
    return re.sub(r'[<>:"/\\|?*]', '_', name) + ext


def translate_text_google(text: str, max_retries=3) -> str:
    """Переклад одного абзацу через Google (deep_translator)."""
    for attempt in range(max_retries):
        try:
            return GoogleTranslator(source='en', target='uk').translate(text)
        except Exception as e:
            logging.warning(f"Google Translator Error (attempt {attempt+1}/{max_retries}): {e}")
            time.sleep(2 ** attempt)
    return "Помилка перекладу (Google)"


# -------------------- OPENAI LOGIC --------------------

def chunk_paragraphs(paragraphs, chunk_size=5):
    """
    Розбиває список абзаців на шматки (chunk) по chunk_size.
    Повертає генератор списків (кожен список — це пакет абзаців).
    """
    for i in range(0, len(paragraphs), chunk_size):
        yield paragraphs[i : i+chunk_size]


def translate_chunk_openai(paragraph_chunk):
    """
    Викликає OpenAI (через OpenRouter) для кількох абзаців за раз.
    """
    if not paragraph_chunk:
        return []

    prompt_text = ""
    for i, para in enumerate(paragraph_chunk, start=1):
        prompt_text += f"{i}) {para}\n"
    
    # Використовуємо модель OpenRouter (з префіксом виробника)
    model_to_use = "openai/gpt-5-mini"
    
    logging.info(f"--- ВІДПРАВКА ЗАПИТУ ДО OPENROUTER ---")
    logging.info(f"Модель: {model_to_use}")

    try:
        response = client.chat.completions.create(
            model=model_to_use,
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert translator. You will receive a numbered list of paragraphs in English. Your task is to translate them into Ukrainian. Your response MUST contain ONLY the translated paragraphs, preserving the original numbering and order. Do not add any extra text, explanations, or greetings."
                },
                {
                    "role": "user",
                    "content": prompt_text
                },
            ],
            extra_headers={
                "HTTP-Referer": "https://legaltransua.streamlit.app/",
                "X-Title": "LegalTransUA"
            }
        )
        result_text = response.choices[0].message.content.strip()

        logging.info(f"--- ОТРИМАНО УСПІШНУ ВІДПОВІДЬ ВІД OPENROUTER ---")
        if not result_text:
            logging.warning("!!! Відповідь ПОРОЖНЯ !!!")

    except Exception as e:
        logging.error(f"!!! ПОМИЛКА ЗАПИТУ ДО OPENROUTER !!!")
        logging.error(f"Повний текст помилки: {e}")
        return ["Помилка перекладу (AI)"] * len(paragraph_chunk)

    # Парсинг відповіді
    lines = result_text.splitlines()
    chunk_result = []
    current_translation = ""
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        match = re.match(r"^\d+\)\s*(.*)", line_stripped)
        
        if match:
            if current_translation:
                chunk_result.append(current_translation.strip())
            current_translation = match.group(1)
        else:
            current_translation += " " + line_stripped

    if current_translation:
        chunk_result.append(current_translation.strip())
        
    logging.info(f"Після розбору знайдено {len(chunk_result)} перекладених абзаців.")

    if not chunk_result and result_text:
        logging.warning("Парсинг не дав результату, але текст існує. Спроба повернути сирий текст.")
        chunk_result.append(result_text)

    while len(chunk_result) < len(paragraph_chunk):
        chunk_result.append("Помилка: не вдалося розпарсити відповідь GPT")

    return chunk_result[: len(paragraph_chunk)]


# -------------------- DOCX FORMATTING --------------------

def setup_document_orientation(doc: Document):
    """Налаштовує горизонтальну орієнтацію документа."""
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def add_title(doc: Document):
    """Додає заголовок у документ."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Документ створено за допомогою скрипта перекладу LegalTransUA")
    run.bold = True
    run.font.size = Pt(12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_shading_element(color: str):
    """Створює елемент заливки комірки в таблиці."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    return shading


def create_translation_table(doc: Document, paragraphs, google_trans, openai_trans):
    """Створює таблицю порівняння."""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["№", "Оригінальний текст", "Google Translate", "OpenAI GPT"]
    header_fill_color = "D9EAF7"
    row_number_fill_color = "E0E0E0"

    # Заголовки
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell._element.get_or_add_tcPr().append(create_shading_element(header_fill_color))

    # Рядки
    for i, (para, g, o) in enumerate(zip(paragraphs, google_trans, openai_trans)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        
        # БЕЗПЕЧНИЙ ЗАПИС: якщо перекладу немає, ставимо порожній рядок
        row_cells[1].text = str(para) if para else ""
        row_cells[2].text = str(g) if g else "" 
        row_cells[3].text = str(o) if o else ""

        row_cells[0]._element.get_or_add_tcPr().append(create_shading_element(row_number_fill_color))

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Ширина колонок
    total_width = Inches(10)
    first_col_width = total_width * 0.05
    other_width = (total_width - first_col_width) / 3.0

    col_widths = [first_col_width, other_width, other_width, other_width]

    for i, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = col_widths[i]

    return doc


def save_translation_document(source, paragraphs, google_trans, openai_trans):
    """Зберігає документ."""
    doc = Document()
    setup_document_orientation(doc)
    add_title(doc)
    doc.add_paragraph(f"Дата та час перекладу: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    create_translation_table(doc, paragraphs, google_trans, openai_trans)

    timestamp_str = datetime.now().strftime("%Y-%m-%d")
    if source.startswith("http"):
        base_name = "Document From Internet"
    else:
        base_name = os.path.splitext(os.path.basename(source))[0]
        base_name = sanitize_filename(base_name)

    output_filename = f"{base_name} (Translated by LTUA {timestamp_str}).docx"

    save_dir = "output"
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    output_path = os.path.join(save_dir, output_filename)
    doc.save(output_path)
    logging.info(f"Документ збережено: {output_path}")

    return output_path


def process_document(source: str, openai_chunk_size=5):
    """Основна логіка обробки."""
    paragraphs = extract_text(source)
    if not paragraphs:
        logging.error("Документ не містить тексту або текст не вдалося витягти.")
        return None

    logging.info(f"Знайдено абзаців: {len(paragraphs)}")

    # GOOGLE
    google_translations = [""] * len(paragraphs)
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(translate_text_google, p): i for i, p in enumerate(paragraphs)}
        for future in tqdm(as_completed(futures), total=len(futures), desc="Google"):
            idx = futures[future]
            google_translations[idx] = future.result()

    # OPENAI (OpenRouter)
    openai_translations = [""] * len(paragraphs)
    all_chunks = list(chunk_paragraphs(paragraphs, chunk_size=openai_chunk_size))
    
    processed_count = 0
    total_paras = len(paragraphs)

    for chunk_index, chunk in enumerate(all_chunks):
        chunk_result = translate_chunk_openai(chunk)
        
        start_idx = chunk_index * openai_chunk_size
        for i, translation in enumerate(chunk_result):
            if start_idx + i < len(openai_translations):
                openai_translations[start_idx + i] = translation

        processed_count += len(chunk)
        logging.info(f"OpenAI chunk {chunk_index+1}/{len(all_chunks)} готово.")

    output_file = save_translation_document(
        source,
        paragraphs,
        google_translations,
        openai_translations
    )
    logging.info(f"Успішно збережено: {output_file}")
    return output_file


if __name__ == "__main__":
    # Для локального тестування
    source_path = input("Введіть URL або шлях до PDF/DOCX-файлу: ").strip()
    process_document(source_path, openai_chunk_size=5)
